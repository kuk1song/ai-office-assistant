import fitz  # PyMuPDF
import docx
import os
import shutil
from typing import Dict, List, Union

# Try to import OCR libraries (optional dependency)
try:
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

def parse_document(file_path: str) -> Dict[str, Union[str, List[str]]]:
    """
    Parses a document (PDF, DOCX, TXT) and extracts its text and images.

    Args:
        file_path: The path to the document file.

    Returns:
        A dictionary with 'text' and 'image_paths' keys.
        Returns a dictionary with an 'error' key if an error occurs.
    """
    _, file_extension = os.path.splitext(file_path)
    file_extension = file_extension.lower()

    if file_extension == '.pdf':
        return _parse_pdf(file_path)
    elif file_extension == '.docx':
        return _parse_docx(file_path)
    elif file_extension == '.txt':
        return _parse_txt(file_path)
    else:
        return {"error": f"Unsupported file format '{file_extension}'. Please provide a .pdf, .docx, or .txt file."}

def _extract_text_from_image(image_path: str) -> str:
    """Extract text from an image using OCR."""
    if not OCR_AVAILABLE:
        return ""
    
    try:
        image = Image.open(image_path)
        text = pytesseract.image_to_string(image)
        return text.strip()
    except Exception as e:
        print(f"Warning: Could not extract text from {image_path}. Reason: {e}")
        return ""

def _parse_pdf(file_path: str) -> Dict[str, Union[str, List[str]]]:
    """Extracts text and images from a PDF file."""
    temp_dir = "temp_images"
    if os.path.exists(temp_dir):
        shutil.rmtree(temp_dir)
    os.makedirs(temp_dir)
    
    text = ""
    image_paths = []
    try:
        doc = fitz.open(file_path)
        for page_num, page in enumerate(doc):
            # Extract text
            page_text = page.get_text()
            text += page_text + "\n\n"

            # Extract tables
            tables = page.find_tables()
            if tables:
                for i, table in enumerate(tables):
                    try:
                        table_data = table.extract()
                        if table_data:
                            markdown_table = "### Table Data\n\n"
                            if len(table_data) > 1:
                                header = table_data[0]
                                markdown_table += "| " + " | ".join(map(str, header)) + " |\n"
                                markdown_table += "| " + " | ".join(["---"] * len(header)) + " |\n"
                                for row in table_data[1:]:
                                    markdown_table += "| " + " | ".join(map(str, row)) + " |\n"
                            else: # Handle single row tables or simple lists
                                for row in table_data:
                                    markdown_table += "| " + " | ".join(map(str, row)) + " |\n"
                            
                            text += markdown_table + "\n\n"
                    except Exception as e:
                        print(f"Warning: Could not extract table {i} on page {page_num+1}. Reason: {e}")

            # Extract images
            image_list = page.get_images(full=True)
            for img_index, img in enumerate(image_list):
                xref = img[0]
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                
                image_ext = base_image["ext"]
                image_filename = f"image_p{page_num+1}_{img_index+1}.{image_ext}"
                image_path = os.path.join(temp_dir, image_filename)
                
                with open(image_path, "wb") as img_file:
                    img_file.write(image_bytes)
                image_paths.append(image_path)
                
        doc.close()
        
        # If no text was found but images exist, try OCR
        if not text.strip() and image_paths:
            print("No text found in PDF, attempting OCR on images...")
            ocr_text = ""
            
            if OCR_AVAILABLE:
                for image_path in image_paths:
                    extracted_text = _extract_text_from_image(image_path)
                    if extracted_text:
                        ocr_text += f"\n\n--- Text from {os.path.basename(image_path)} ---\n"
                        ocr_text += extracted_text + "\n"
                
                if ocr_text.strip():
                    text = "=== Text extracted from images using OCR ===\n" + ocr_text
                    print(f"Successfully extracted text from {len(image_paths)} image(s) using OCR.")
                else:
                    text = "=== Document contains only images with no readable text ==="
                    print("OCR did not find any readable text in the images.")
            else:
                text = "=== Document contains only images (OCR not available) ==="
                print("Document contains only images, but OCR libraries are not installed.")
        
        return {"text": text, "image_paths": image_paths}
    except Exception as e:
        return {"error": f"Failed to parse PDF file at {file_path}. Reason: {e}"}

def _parse_docx(file_path: str) -> Dict[str, Union[str, List[str]]]:
    """Extracts text, tables, and images from a DOCX file."""
    temp_dir = "temp_images"
    if not os.path.exists(temp_dir):
        os.makedirs(temp_dir)
        
    text = ""
    image_paths = []
    try:
        doc = docx.Document(file_path)
        
        # Extract paragraphs
        for para in doc.paragraphs:
            text += para.text + "\n"
        text += "\n"

        # Extract tables
        for i, table in enumerate(doc.tables):
            text += f"### Table {i+1} Data\n\n"
            # Create header
            header = [cell.text for cell in table.rows[0].cells]
            text += "| " + " | ".join(header) + " |\n"
            text += "| " + " | ".join(["---"] * len(header)) + " |\n"
            # Create rows
            for row in table.rows[1:]:
                row_data = [cell.text for cell in row.cells]
                text += "| " + " | ".join(row_data) + " |\n"
            text += "\n"
        
        # Extract images
        for i, rel in enumerate(doc.part.rels.values()):
            if "image" in rel.target_ref:
                image = rel.target_part.blob
                ext = rel.target_ref.split('.')[-1]
                
                image_filename = f"image_docx_{i+1}.{ext}"
                image_path = os.path.join(temp_dir, image_filename)
                
                with open(image_path, "wb") as img_file:
                    img_file.write(image)
                image_paths.append(image_path)

        return {"text": text, "image_paths": image_paths}
    except Exception as e:
        return {"error": f"Failed to parse DOCX file at {file_path}. Reason: {e}"}

def _parse_txt(file_path: str) -> Dict[str, Union[str, List[str]]]:
    """Extracts text from a TXT file (no images)."""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        return {"text": text, "image_paths": []}
    except Exception as e:
        return {"error": f"Failed to read TXT file at {file_path}. Reason: {e}"} 